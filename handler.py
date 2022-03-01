from os import listdir
from docx import Document
#Document is used to read .docx files and listdir is used to enumerate the content of a dir


def retrieve_text(pth):
    doc=Document(pth)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return ' '.join(fullText)

def retrieve_text2(pth):
    doc=Document(pth)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText

def retrieve_image(pth):
    images=listdir(pth)
    food=''
    front=''
    for i in range(0, len(images)):
        if "food" in images[i]:
            food=pth+"/"+images[i]
        elif "front" in images[i]:
            front=pth+"/"+images[i]
        else:
            images[i]=pth+"/"+images[i]
    return images, front, food

